import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import re
import asyncio
import edge_tts

# --- 1. KONFIGURASI HALAMAN ---
st.set_page_config(
    page_title="Magis AI - Jesuit Order Edition",
    page_icon="🕊️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Inisialisasi Session State
if 'result_text' not in st.session_state:
    st.session_state.result_text = ""
if 'topic_context' not in st.session_state:
    st.session_state.topic_context = ""
if 'full_audio_bytes' not in st.session_state:
    st.session_state.full_audio_bytes = None

# --- 2. FUNGSI AUDIO ---
async def generate_audio_stream(text, voice, rate):
    try:
        communicate = edge_tts.Communicate(text, voice, rate=rate)
        audio_data = b""
        async for chunk in communicate.stream():
            if chunk["type"] == "audio":
                audio_data += chunk["data"]
        return audio_data
    except Exception as e:
        return None

def get_audio(text, lang_code, gender, speed_pct):
    if not text or not text.strip(): return None
    clean_text_audio = re.sub(r'[#*_`\-]', '', text)
    clean_text_audio = re.sub(r'\|', ' ', clean_text_audio)
    
    if lang_code == "ID":
        voice = "id-ID-ArdiNeural" if gender == "Pria" else "id-ID-GadisNeural"
    else:
        voice = "en-US-ChristopherNeural" if gender == "Male" else "en-US-AriaNeural"
    
    rate_str = f"{int(speed_pct)}%"
    if speed_pct >= 0: rate_str = f"+{rate_str}"
    
    try:
        audio_bytes = asyncio.run(generate_audio_stream(clean_text_audio, voice, rate_str))
        return audio_bytes
    except Exception as e:
        st.error(f"Audio Error: {e}")
        return None

# --- 3. KAMUS NILAI & PROMPT ---
TRANS = {
    "ID": {
        "title_sub": "Mitra Diskresi Guru Ignasian",
        "sidebar_settings": "Pengaturan",
        "lbl_lang": "Bahasa",
        "lbl_theme": "Tampilan",
        "lbl_tone": "Gaya Bahasa AI",
        "opt_tone": ["Reflektif (Spirit Jesuit)", "Akademis (St. Petrus Kanisius)", "Pastoral (Cura Personalis)"],
        "lbl_key": "Kunci Akses Google",
        "lbl_model": "Model Kecerdasan",
        "lbl_voice_set": "Pengaturan Suara (Audio)",
        "lbl_gender": "Suara Narator",
        "lbl_speed": "Kecepatan Bicara",
        "lbl_menu": "Pilih Modul Formasi",
        "menu_opt": ["1. Konteks (Cura Personalis & UAP)", "2. Desain RPP (IPP & UAP)", "3. Refleksi (Examen & Magis)", "4. Asisten Makalah (Jesuit Scholar)"],
        "btn_analyze": "Analisis Profil Siswa",
        "btn_rpp": "Desain Pembelajaran PPI",
        "btn_reflect": "Mulai Diskresi",
        "btn_makalah": "Susun Draf Makalah",
        "btn_dl_word": "📥 Unduh Dokumen (.docx)",
        "loading": "✨ Sedang menimbang (Diskresi) dalam terang UAP & Magis...",
        "empty_warning": "⚠️ Mohon isi data untuk memulai proses.",
        "key_warning": "🔒 Kunci Akses (API Key) diperlukan. Masukkan di menu sebelah kiri.",
        "key_missing_alert": "⚠️ **PERHATIAN:** API Key belum dimasukkan. Silakan buka Sidebar (kiri) bagian 'Kunci Akses Google' agar AI dapat bekerja.",
        "core_values": """
        **CORE PARADIGM (THE JESUIT WAY):**
        1. **Universal Apostolic Preferences (UAP):** Menunjukkan Jalan kepada Allah, Berjalan bersama yang Terpinggirkan, Menemani Orang Muda, Merawat Rumah Kita Bersama.
        2. **Values (4C + 1L):** Competence, Conscience, Compassion, Commitment, Leadership.
        3. **Magis:** Selalu mencari yang 'lebih' demi kemuliaan Allah (AMDG).
        """,
        "m1_t": "📘 Konteks (Cura Personalis & UAP)", 
        "m1_l1": "Profil Unik Siswa / Dinamika Kelas:", 
        "m1_p1": "Ceritakan karakter, latar belakang, atau tantangan siswa...", 
        "m1_l2": "Fokus Masalah / Situasi:", 
        "m1_sys": "ROLE: Pendidik Ignasian (Cura Personalis). Analisis siswa dengan lensa 4C+1L dan UAP.",
        "m2_t": "📙 Desain Pembelajaran (IPP & UAP)", 
        "m2_l1": "Topik / Materi Pembelajaran:", 
        "m2_l2": "Durasi & Target Formasi:", 
        "m2_sys": "ROLE: Perancang IPP (Ignatian Pedagogical Paradigm). Siklus: Context-Experience-Reflection-Action-Evaluation.",
        "m3_t": "📗 Refleksi Batin (Examen)", 
        "m3_l1": "Peristiwa / Kegelisahan / Topik Refleksi:", 
        "m3_sys": "ROLE: Pembimbing Rohani (Examen). Pandu 5 langkah (Syukur, Mohon Terang, Tinjauan, Sesal, Niat).",
        "m4_t": "📘 Asisten Makalah (Jesuit Scholar)", 
        "m4_l1": "Topik / Judul Makalah:", 
        "m4_l2": "Argumen Utama (Tesis):",
        "m4_p2": "Contoh: Teknologi AI harus diarahkan untuk memuliakan Tuhan (AMDG)...",
        "m4_cat_lbl": "Pilih Kedalaman Analisis:",
        "m4_cat_opt": ["Ringkas (Poin Utama 4C)", "Standar (Esai Akademis)", "Mendalam (Diskresi Teologis & UAP)"],
        "m4_prompt_ringkas": "STYLE: To the point, Bullet Points. FOCUS: Identifikasi elemen Competence, Conscience, Compassion.",
        "m4_prompt_standar": "STYLE: Akademis, Naratif, Argumentatif (Style St. Petrus Kanisius). FOCUS: Bangun argumen logis + moral.",
        "m4_prompt_mendalam": "STYLE: Filosofis, Teologis, Reflektif Mendalam (Style St. Ignatius). FOCUS: Diskresi mendalam, Tantum Quantum.",
    },
    "EN": {
        "title_sub": "Ignatian Teacher's Discernment Partner",
        "sidebar_settings": "Settings",
        "lbl_lang": "Language",
        "lbl_theme": "Theme",
        "lbl_tone": "AI Tone",
        "opt_tone": ["Reflective (Jesuit Spirit)", "Academic (Canisius)", "Pastoral (Cura Personalis)"],
        "lbl_key": "Google Access Key",
        "lbl_model": "AI Model",
        "lbl_voice_set": "Audio Settings",
        "lbl_gender": "Narrator Voice",
        "lbl_speed": "Speech Speed",
        "lbl_menu": "Select Formation Module",
        "menu_opt": ["1. Context (Cura Personalis)", "2. Lesson Design (IPP & UAP)", "3. Reflection (Examen)", "4. Paper Assistant (Jesuit Scholar)"],
        "btn_analyze": "Analyze Profile",
        "btn_rpp": "Design Lesson",
        "btn_reflect": "Start Discernment",
        "btn_makalah": "Draft Paper",
        "btn_dl_word": "📥 Download Word",
        "loading": "✨ Discerning with Magis & UAP Spirit...",
        "empty_warning": "⚠️ Please provide input.",
        "key_warning": "🔒 API Key required. Please check sidebar.",
        "key_missing_alert": "⚠️ **ATTENTION:** API Key is missing. Please enter it in the Sidebar (left) to proceed.",
        "core_values": """
        **CORE PARADIGM (THE JESUIT WAY):**
        1. **UAP:** Showing the Way to God, Walking with the Excluded, Journeying with Youth, Caring for our Common Home.
        2. **Values:** Competence, Conscience, Compassion, Commitment, Leadership.
        3. **Magis:** AMDG.
        """,
        "m1_t": "📘 Context (Cura Personalis)", 
        "m1_l1": "Student Profile:", 
        "m1_p1": "Describe character...", 
        "m1_l2": "Situation:", 
        "m1_sys": "ROLE: Canisian Educator. Analyze via 4C+1L & UAP.",
        "m2_t": "📙 Lesson Design (IPP)", 
        "m2_l1": "Topic:", 
        "m2_l2": "Goal:", 
        "m2_sys": "ROLE: IPP Designer. Cycle: Context-Experience-Reflection-Action-Evaluation.",
        "m3_t": "📗 Reflection (Examen)", 
        "m3_l1": "Topic:", 
        "m3_sys": "ROLE: Spiritual Director. Guide Examen.",
        "m4_t": "📘 Academic Paper Assistant", 
        "m4_l1": "Paper Topic:", 
        "m4_l2": "Thesis:",
        "m4_p2": "E.g., AI technology must be directed to glorify God...",
        "m4_cat_lbl": "Select Depth:",
        "m4_cat_opt": ["Concise (4C Points)", "Standard (Academic Essay)", "Deep (Theological Discernment)"],
        "m4_prompt_ringkas": "STYLE: Bullet points. FOCUS: Key 4C elements.",
        "m4_prompt_standar": "STYLE: Academic (Canisius style). FOCUS: Logical arguments + values.",
        "m4_prompt_mendalam": "STYLE: Philosophical/Theological (Ignatian). FOCUS: Deep discernment.",
    }
}

# --- 4. LOGIKA WORD PROCESSING ---
def clean_text(text):
    text = text.replace('**', '').replace('__', '')
    text = text.replace('```', '')
    text = re.sub(r'\$(.*?)\$', r'\1', text) 
    return text

def process_markdown_to_docx(doc, text):
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    lines = text.split('\n')
    table_buffer = [] 
    
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('|') and stripped.endswith('|'):
            cells = [c.strip() for c in stripped.split('|')]
            if len(cells) > 2: cells = cells[1:-1] 
            if '---' in cells[0]: continue 
            table_buffer.append(cells)
        else:
            if table_buffer:
                rows = len(table_buffer)
                cols = len(table_buffer[0])
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                for i, row_data in enumerate(table_buffer):
                    for j, cell_text in enumerate(row_data):
                        if j < len(table.rows[i].cells):
                            cell_p = table.rows[i].cells[j].paragraphs[0]
                            run = cell_p.add_run(clean_text(cell_text))
                            run.font.name = 'Arial'
                            run.font.size = Pt(12)
                            cell_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT 
                doc.add_paragraph() 
                table_buffer = [] 

            if stripped:
                clean_line = clean_text(stripped)
                if stripped.startswith('### '):
                    h = doc.add_heading(clean_line.replace('### ', ''), level=2)
                    h.runs[0].font.name = 'Arial'
                    h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif stripped.startswith('## '):
                    h = doc.add_heading(clean_line.replace('## ', ''), level=1)
                    h.runs[0].font.name = 'Arial'
                    h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif stripped.startswith('- ') or stripped.startswith('* '):
                    p = doc.add_paragraph(clean_line.replace('- ', '').replace('* ', ''), style='List Bullet')
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    if p.runs: p.runs[0].font.name = 'Arial'
                elif re.match(r'^\d+\.', stripped):
                    p = doc.add_paragraph(re.sub(r'^\d+\.\s', '', clean_line), style='List Number')
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    if p.runs: p.runs[0].font.name = 'Arial'
                else:
                    p = doc.add_paragraph(clean_line)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    if p.runs: p.runs[0].font.name = 'Arial'
    
    if table_buffer: # Flush sisa tabel
        rows = len(table_buffer)
        cols = len(table_buffer[0])
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        for i, row_data in enumerate(table_buffer):
            for j, cell_text in enumerate(row_data):
                if j < len(table.rows[i].cells):
                    cell_p = table.rows[i].cells[j].paragraphs[0]
                    run = cell_p.add_run(clean_text(cell_text))
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)

def create_docx(content, topic, lang_key):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    header_text = 'MAGIS AI - JESUIT ORDER RESULT'
    h = doc.add_heading(header_text, 0)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in h.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    runner = p.add_run(f'Topic: {topic}')
    runner.bold = True
    runner.font.name = 'Arial'
    runner.font.size = Pt(12)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('\n')
    
    process_markdown_to_docx(doc, content)
    
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_text = 'Dicetak oleh Magis AI - Kolese Kanisius (A. Henny Setyawan)'
    f_run = footer_para.add_run(f'\n--- {footer_text} ---')
    f_run.font.name = 'Arial'
    f_run.font.size = Pt(9)
    f_run.italic = True
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    bio = BytesIO()
    doc.save(bio)
    return bio

def get_gemini_response(api_key, model_name, system_instruction, user_prompt, tone, lang, core_vals):
    if not api_key: return "⚠️ Error: API Key Missing"
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        lang_instruction = "Use standard Indonesian." if lang == "ID" else "Use professional English."
        # Meminta format tabel secara eksplisit
        format_instruction = "FORMAT: Use Markdown Tables for structured data clearly. Use standard paragraphs."
        full_sys = f"ROLE: {system_instruction}\n\n{core_vals}\n\nTONE: {tone}\nLANGUAGE: {lang_instruction}\n{format_instruction}"
        
        response = model.generate_content(f"{full_sys}\n\nTASK: {user_prompt}")
        text = response.text
        if text.startswith("```"): text = text.replace("```markdown", "").replace("```", "")
        return text
    except Exception as e: return f"Error: {str(e)}"

# --- 5. CSS (THE "MASTERPIECE" THEME) ---
def inject_custom_css(theme):
    if theme == "Gelap":
        # DARK MODE (Elegant & Deep)
        bg_color = "#0E1117"
        sidebar_bg = "#161B22"
        text_color = "#E6EDF3"        # Putih Tulang (Nyaman)
        secondary_text = "#B0B8C4"
        input_bg = "#0d1117"
        input_border = "#30363D"
        card_bg = "#161B22"
        table_border = "#30363D"
        table_header_bg = "rgba(66, 133, 244, 0.2)"
        banner_overlay = "rgba(0, 0, 0, 0.6)"
    else:
        # LIGHT MODE (Soft & High Contrast)
        bg_color = "#F3F4F6"          # Soft Porcelain (Abu sangat muda, tidak putih buta)
        sidebar_bg = "#FFFFFF"        # Putih Bersih
        text_color = "#111827"        # Hitam Pekat (Charcoal) - SANGAT TERBACA
        secondary_text = "#4B5563"
        input_bg = "#FFFFFF"
        input_border = "#D1D5DB"
        card_bg = "#FFFFFF"
        table_border = "#E5E7EB"
        table_header_bg = "#F9FAFB"
        banner_overlay = "rgba(0, 0, 0, 0.4)"

    st.markdown(f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');
    
    :root {{ 
        --bg-color: {bg_color}; 
        --sidebar-bg: {sidebar_bg};
        --text-color: {text_color};
        --secondary-text: {secondary_text};
        --input-bg: {input_bg};
        --input-border: {input_border};
        --card-bg: {card_bg};
        --table-border: {table_border};
        --table-header-bg: {table_header_bg};
        --banner-overlay: {banner_overlay};
        --primary-color: #2563EB;
    }}
    
    /* 1. RESET & BASE */
    html, body, .stApp {{ 
        background-color: var(--bg-color) !important; 
        color: var(--text-color) !important;
        font-family: 'Inter', sans-serif; 
    }}

    /* 2. TEXT VISIBILITY ENFORCEMENT (AGRESSIVE) */
    h1, h2, h3, h4, h5, h6, p, li, label, span, div[data-testid="stMarkdownContainer"] p {{
        color: var(--text-color) !important;
    }}

    /* 3. TOMBOL (EXCEPTION) - Tetap Putih & Gradasi */
    div.stButton > button {{ 
        background: linear-gradient(135deg, #2563EB, #10B981) !important; 
        color: #FFFFFF !important; 
        border: none; 
        border-radius: 50px; 
        padding: 0.6rem 1.8rem; /* Lebih besar sedikit biar enak ditekan */
        font-weight: 700; 
        letter-spacing: 0.5px;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1); 
        transition: all 0.2s; 
    }}
    div.stButton > button p {{
        color: #FFFFFF !important;
    }}
    div.stButton > button:hover {{ 
        transform: scale(1.02); 
        box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.2); 
    }}

    /* 4. SIDEBAR */
    section[data-testid="stSidebar"] {{ 
        background-color: var(--sidebar-bg) !important; 
        border-right: 1px solid var(--input-border);
    }}
    section[data-testid="stSidebar"] * {{
        color: var(--text-color) !important;
    }}

    /* 5. INPUT FIELDS */
    .stTextInput input, .stTextArea textarea, .stSelectbox div[data-baseweb="select"] > div {{ 
        background-color: var(--input-bg) !important; 
        color: var(--text-color) !important; 
        border: 1px solid var(--input-border) !important; 
        border-radius: 8px; 
    }}

    /* 6. RESULT CARD (RATA KANAN KIRI) */
    .result-card {{ 
        background-color: var(--card-bg); 
        border: 1px solid var(--input-border); 
        border-radius: 12px; 
        padding: 40px; /* Padding luas agar lega */
        box-shadow: 0 2px 5px 0 rgba(0, 0, 0, 0.05); 
        text-align: justify !important; 
        line-height: 1.8; /* Jarak antar baris enak dibaca */
    }}
    .result-card p {{
        text-align: justify !important;
        margin-bottom: 1em;
    }}

    /* 7. TABEL RAPI (STYLED TABLES) */
    .result-card table {{
        width: 100%;
        border-collapse: collapse;
        margin: 20px 0;
        font-size: 0.95em;
        border-radius: 8px;
        overflow: hidden; /* Biar border radius ngaruh ke header */
        border: 1px solid var(--table-border);
    }}
    .result-card th {{
        background-color: var(--table-header-bg);
        color: var(--text-color) !important;
        font-weight: bold;
        padding: 12px 15px;
        text-align: left;
        border-bottom: 2px solid var(--table-border);
    }}
    .result-card td {{
        padding: 12px 15px;
        border-bottom: 1px solid var(--table-border);
        color: var(--text-color) !important;
    }}
    .result-card tr:last-of-type td {{
        border-bottom: none;
    }}

    /* 8. HEADER BANNER */
    .main-title-text {{
        font-weight: 900;
        font-size: 3.8rem;
        color: #FFFFFF !important; 
        text-shadow: 0 2px 15px rgba(0,0,0,0.8); 
    }}
    .subtitle-text {{
        color: rgba(255, 255, 255, 0.95) !important;
        font-size: 1.2rem;
        font-weight: 500;
        text-shadow: 0 1px 5px rgba(0,0,0,0.8);
    }}
    .title-container {{
        position: relative;
        overflow: hidden;
        padding: 3.5rem 1rem; 
        margin-bottom: 1.5rem;
        border-radius: 16px;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
    }}
    .title-container::before {{
        content: "";
        position: absolute;
        top: 0; left: 0; width: 100%; height: 100%;
        background-image: 
            linear-gradient(var(--banner-overlay), var(--banner-overlay)),
            url('https://i.imgur.com/wmAE0d7.jpeg');
        background-size: cover;
        background-position: center 30%; 
        z-index: 0;
    }}
    .title-content {{
        position: relative;
        z-index: 1;
        text-align: center;
    }}
    
    #MainMenu, footer {{visibility: hidden;}}
    [data-testid="stDeployButton"] {{display: none;}}
    </style>
    """, unsafe_allow_html=True)

# --- 6. TAMPILAN UI ---
with st.sidebar:
    st.markdown("<div style='text-align:center; margin-bottom:20px;'><img src='https://i.imgur.com/UUCgyfV.png' width='90'></div>", unsafe_allow_html=True)
    lang_opt = st.radio("Bahasa", ["Indonesia 🇮🇩", "English 🇺🇸"], horizontal=True, label_visibility="collapsed")
    L_CODE = "ID" if "Indonesia" in lang_opt else "EN"
    TXT = TRANS[L_CODE] 
    
    st.markdown(f"### ⚙️ {TXT['sidebar_settings']}")
    theme_opt = st.radio(TXT["lbl_theme"], ["Modern Dark 🌑", "Clean Light ☀️"])
    THEME_VAL = "Gelap" if "Dark" in theme_opt else "Terang"
    
    tone_idx = st.selectbox(TXT["lbl_tone"], TXT["opt_tone"])
    st.divider()
    
    # --- AUDIO SETTINGS ---
    st.markdown(f"### 🔊 {TXT['lbl_voice_set']}")
    gender_opt = ["Pria", "Wanita"] if L_CODE == "ID" else ["Male", "Female"]
    sel_gender = st.radio(TXT['lbl_gender'], gender_opt, horizontal=True)
    sel_speed = st.slider(TXT['lbl_speed'], -50, 50, 0, format="%d%%")
    st.divider()
    
    api_key = st.text_input(TXT["lbl_key"], type="password")
    models = ["gemini-pro", "gemini-1.5-flash"] 
    if api_key:
        try:
            genai.configure(api_key=api_key)
            fetched_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
            if fetched_models: models = fetched_models
        except: pass
    sel_model = st.selectbox(TXT["lbl_model"], models)
    
    st.divider()
    menu_sel = st.radio(TXT["lbl_menu"], TXT["menu_opt"])
    menu_idx = TXT["menu_opt"].index(menu_sel)
    
    # CREDIT TITLE
    st.markdown(f"""
    <div style='margin-top:3rem;text-align:center;font-size:0.75rem;opacity:0.8;line-height:1.6;border-top:1px solid var(--input-border); padding-top:1rem;'>
    <strong>MAGIS AI v16.0</strong><br>
    Jesuit Order & UAP Integrated<br>
    Design by: <strong>Albertus Henny Setyawan</strong><br>
    Kolese Kanisius Jakarta | 2026
    </div>
    """, unsafe_allow_html=True)

# INJECT CSS AFTER THEME SELECTION
inject_custom_css(THEME_VAL)

# MAIN CONTENT
st.markdown(f"""
<div class='title-container'>
    <div class='title-content'>
        <h1 class='main-title-text'>MAGIS AI</h1>
        <div class='subtitle-text'>{TXT['title_sub']}</div>
    </div>
</div>
""", unsafe_allow_html=True)

# --- PERINGATAN API KEY ---
if not api_key:
    st.warning(TXT['key_missing_alert'])

with st.container():
    prompt = ""
    sys_instruction = ""
    execute = False 
    
    if menu_idx == 0: 
        st.markdown(f"<h3 style='color:#4285F4 !important;'>{TXT['m1_t']}</h3>", unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        in_1 = c1.text_area(TXT['m1_l1'], placeholder=TXT['m1_p1'], height=150)
        in_2 = c2.text_input(TXT['m1_l2'])
        if st.button(TXT['btn_analyze']):
            execute = True
            if in_1:
                prompt = f"{TXT['m1_l1']} {in_1} | {TXT['m1_l2']} {in_2}"
                sys_instruction = TXT['m1_sys']

    elif menu_idx == 1: 
        st.markdown(f"<h3 style='color:#F59E0B !important;'>{TXT['m2_t']}</h3>", unsafe_allow_html=True)
        c1, c2 = st.columns([2, 1])
        in_1 = c1.text_input(TXT['m2_l1'])
        in_2 = c2.selectbox(TXT['m2_l2'], ["1 JP (45')", "2 JP (90')", "Block Project (Project Based)"])
        if st.button(TXT['btn_rpp']):
            execute = True
            if in_1:
                prompt = f"Topik: {in_1} | Durasi: {in_2}"
                sys_instruction = TXT['m2_sys']

    elif menu_idx == 2:
        st.markdown(f"<h3 style='color:#10B981 !important;'>{TXT['m3_t']}</h3>", unsafe_allow_html=True)
        in_1 = st.text_area(TXT['m3_l1'], height=100)
        if st.button(TXT['btn_reflect']):
            execute = True
            if in_1:
                prompt = f"Bahan Refleksi: {in_1}"
                sys_instruction = TXT['m3_sys']

    elif menu_idx == 3: 
        st.markdown(f"<h3 style='color:#4285F4 !important;'>{TXT['m4_t']}</h3>", unsafe_allow_html=True)
        in_1 = st.text_input(TXT['m4_l1'])
        in_2 = st.text_area(TXT['m4_l2'], placeholder=TXT['m4_p2'], height=100)
        cat_sel = st.radio(TXT['m4_cat_lbl'], TXT['m4_cat_opt'], horizontal=True)
        
        if st.button(TXT['btn_makalah']):
            execute = True
            if in_1 and in_2:
                cat_prompt = ""
                if "Ringkas" in cat_sel or "Concise" in cat_sel:
                    cat_prompt = TXT['m4_prompt_ringkas']
                elif "Standar" in cat_sel or "Standard" in cat_sel:
                    cat_prompt = TXT['m4_prompt_standar']
                else:
                    cat_prompt = TXT['m4_prompt_mendalam']
                
                prompt = f"TOPIK: {in_1}\nTESIS: {in_2}"
                sys_instruction = f"""
                PERAN: Cendekiawan Jesuit.
                KATEGORI: {cat_sel}.
                {cat_prompt}
                TUGAS: Susun makalah sesuai kategori.
                """

    # --- EKSEKUSI AI ---
    if execute:
        if not api_key:
            st.error(TXT["key_warning"]) 
        elif prompt and sys_instruction:
            with st.spinner(TXT['loading']):
                res = get_gemini_response(api_key, sel_model, sys_instruction, prompt, tone_idx, L_CODE, TXT['core_values'])
                st.session_state.result_text = res
                st.session_state.topic_context = prompt
                
                # Full Audio
                full_audio = get_audio(res, L_CODE, sel_gender, sel_speed) 
                st.session_state.full_audio_bytes = full_audio

# --- OUTPUT AREA ---
if st.session_state.result_text:
    st.markdown("---")
    
    c_res, c_audio = st.columns([3.5, 1])
    
    with c_res:
        st.markdown(f"<div class='result-card'>{st.session_state.result_text}</div>", unsafe_allow_html=True)
    
    with c_audio:
        st.markdown(f"<div style='background-color:var(--card-bg); padding:15px; border-radius:10px; border:1px solid var(--input-border); text-align:center;'>", unsafe_allow_html=True)
        st.markdown(f"#### 🎧 Audio")
        if st.session_state.full_audio_bytes:
            st.audio(st.session_state.full_audio_bytes, format='audio/mp3')
        else:
            st.caption("Audio unavailable")
        st.markdown("</div>", unsafe_allow_html=True)
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        docx_file = create_docx(st.session_state.result_text, st.session_state.topic_context, L_CODE)
        st.download_button(
            label=TXT['btn_dl_word'],
            data=docx_file.getvalue(),
            file_name=f"MagisAI_{L_CODE}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )